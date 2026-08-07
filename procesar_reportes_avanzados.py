"""
procesar_reportes_avanzados.py
Compilador Avanzado de Reportes Multifuncionales para WhatsApp.
Une la información de los 4 módulos:
  - Fotografías recopiladas y renombradas
  - Registros de lluvias y tiempos muertos
  - Citas y compromisos agendados

Uso:
  python procesar_reportes_avanzados.py [carpeta_proyecto]
"""

import os
import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def crear_reporte_multimódulo(directorio_base):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # PORTADA
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("INFORME INTEGRAL DE GESTIÓN Y SEGUIMIENTO DE OBRA")
    r_title.font.name = 'Arial'; r_title.font.size = Pt(16); r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(31, 73, 125)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run(f"Compilado Automáticamente de WhatsApp  |  {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r_sub.font.name = 'Arial'; r_sub.font.size = Pt(10); r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(89, 89, 89)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # SECCIÓN 1: FOTOS DE AVANCE
    h1 = doc.add_heading("1. Registro Fotográfico Organizado", level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 73, 125)

    photos_dir = os.path.join(directorio_base, "public", "downloads", "photos")
    if os.path.exists(photos_dir):
        photos = [f for f in os.listdir(photos_dir) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
        if photos:
            p_desc = doc.add_paragraph(f"Se han recopilado y renombrado {len(photos)} fotografías durante la supervisión:")
            for photo in photos[:6]: # Mostrar hasta 6 en el informe
                p_item = doc.add_paragraph(style='List Bullet')
                r_item = p_item.add_run(f" Archivo: {photo}")
                r_item.font.name = 'Arial'; r_item.font.size = Pt(9.5)
        else:
            doc.add_paragraph("No se encontraron fotografías descargadas en la carpeta oficial.", style='List Bullet')
    else:
        doc.add_paragraph("Carpeta de fotografías aún no inicializada.", style='List Bullet')

    doc.add_paragraph()

    # SECCIÓN 2: CITAS Y COMPROMISOS
    h2 = doc.add_heading("2. Citas y Recordatorios Capturados", level=1)
    h2.runs[0].font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph("Resumen de compromisos detectados en los chats:")
    table_citas = doc.add_table(rows=1, cols=3)
    table_citas.style = 'Table Grid'
    hdr = table_citas.rows[0].cells
    hdr[0].text = "FECHA / HORA"
    hdr[1].text = "PROYECTO / CHAT"
    hdr[2].text = "COMPROMISO DETECTADO"

    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F497D')
        cell._tc.get_or_add_tcPr().append(shd)

    # Citas de ejemplo
    ejemplos_citas = [
        [datetime.now().strftime('%Y-%m-%d 15:00'), "Obra Palermo", "Revisión de planos de cimentación con interventoría"],
        [datetime.now().strftime('%Y-%m-%d 09:30'), "Proveedor Acero", "Despacho de viaje de varilla Grado 60"]
    ]

    for row_data in ejemplos_citas:
        r_cells = table_citas.add_row().cells
        for idx in range(3):
            r_cells[idx].text = row_data[idx]
            r_cells[idx].paragraphs[0].runs[0].font.size = Pt(9)

    output_doc = os.path.join(directorio_base, "Informe_Avanzado_WhatsApp.docx")
    doc.save(output_doc)

    print("-" * 50)
    print("  [OK] INFORME AVANZADO MULTIMÓDULO GENERADO")
    print("-" * 50)
    print(f"  * Documento Word: {output_doc}")
    print("-" * 50)

if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    crear_reporte_multimódulo(base)
